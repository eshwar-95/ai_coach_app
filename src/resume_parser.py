"""Resume and document parsing module."""
import io
import tempfile
from typing import Optional
import PyPDF2
from docx import Document
from src.config import ALLOWED_RESUME_FORMATS, MAX_RESUME_SIZE


class ResumeParser:
    """Parser for various resume formats."""

    @staticmethod
    def validate_file(uploaded_file) -> bool:
        """
        Validate uploaded resume file.

        Args:
            uploaded_file: Streamlit uploaded file object

        Returns:
            True if file is valid, False otherwise
        """
        if uploaded_file is None:
            return False

        # Check file size
        if uploaded_file.size > MAX_RESUME_SIZE:
            return False

        # Check file extension
        file_ext = uploaded_file.name.split(".")[-1].lower()
        return file_ext in ALLOWED_RESUME_FORMATS

    @staticmethod
    def parse_resume(uploaded_file) -> Optional[str]:
        """
        Parse resume from uploaded file.

        Args:
            uploaded_file: Streamlit uploaded file object

        Returns:
            Extracted text from resume or None if parsing fails
        """
        if not ResumeParser.validate_file(uploaded_file):
            return None

        file_ext = uploaded_file.name.split(".")[-1].lower()

        try:
            if file_ext == "txt":
                return ResumeParser._parse_txt(uploaded_file)
            elif file_ext == "pdf":
                return ResumeParser._parse_pdf(uploaded_file)
            elif file_ext in ["doc", "docx"]:
                return ResumeParser._parse_docx(uploaded_file)
        except Exception as e:
            raise ValueError(f"Error parsing resume: {str(e)}")

        return None

    @staticmethod
    def _parse_txt(uploaded_file) -> str:
        """Parse text file."""
        return uploaded_file.read().decode("utf-8", errors="ignore")

    @staticmethod
    def _parse_pdf(uploaded_file) -> str:
        """Parse PDF file."""
        text = ""
        try:
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(uploaded_file.read()))
            for page in pdf_reader.pages:
                text += page.extract_text()
        except Exception as e:
            raise ValueError(f"Error parsing PDF: {str(e)}")
        return text

    @staticmethod
    def _parse_docx(uploaded_file) -> str:
        """Parse DOCX/DOC file."""
        text = ""
        try:
            doc = Document(io.BytesIO(uploaded_file.read()))
            for para in doc.paragraphs:
                text += para.text + "\n"
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
        except Exception as e:
            raise ValueError(f"Error parsing DOCX: {str(e)}")
        return text

    @staticmethod
    def extract_skills_from_text(text: str) -> list:
        """
        Extract potential skills from resume text.
        This is a basic implementation - can be enhanced with NLP.

        Args:
            text: Resume text

        Returns:
            List of potential skills
        """
        # Common tech skills to look for
        common_skills = [
            "python",
            "java",
            "javascript",
            "typescript",
            "sql",
            "react",
            "angular",
            "vue",
            "nodejs",
            "django",
            "flask",
            "spring",
            "aws",
            "azure",
            "gcp",
            "docker",
            "kubernetes",
            "git",
            "ci/cd",
            "ml",
            "ai",
            "data science",
            "machine learning",
            "deep learning",
            "tensorflow",
            "pytorch",
            "pandas",
            "numpy",
            "power bi",
            "tableau",
            "excel",
            "salesforce",
            "sap",
            "linux",
            "windows",
            "agile",
            "scrum",
            "jira",
        ]

        text_lower = text.lower()
        found_skills = [skill for skill in common_skills if skill in text_lower]
        return list(set(found_skills))
